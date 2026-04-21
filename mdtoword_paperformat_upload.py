#pip install pypandoc python-docx
#上标：{{SUP_START}}内容{{SUP_END}}
import pypandoc
import os
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def set_format(paragraph, font_east, font_west, size, is_bold=False, align=None, is_title=False):
    if align:
        paragraph.alignment = align
    paragraph.paragraph_format.line_spacing = 1.5 if is_title else 1.25

    # 合并文本以处理跨 run 的暗号
    full_text = "".join(run.text for run in paragraph.runs)
    
    # --- 核心修改：匹配 {{SUP_START}} 和 {{SUP_END}} ---
    if '{{SUP_START}}' in full_text:
        # 使用正则切分出所有的暗号块
        parts = re.split(r'({{SUP_START}}.*?{{SUP_END}})', full_text)
        paragraph.clear() # 清空旧内容，咱们按新规矩重组
        
        for part in parts:
            new_run = paragraph.add_run()
            # 判定当前块是否为暗号内容
            if part.startswith('{{SUP_START}}') and part.endswith('{{SUP_END}}'):
                # 关键点：{{SUP_START}} 占 13 位，{{SUP_END}} 占 11 位
                new_run.text = part[13:-11] 
                new_run.font.superscript = True
                
                # 强力注入 XML 指令，模拟 Ctrl+Shift++
                rPr = new_run._element.get_or_add_rPr()
                from docx.oxml import OxmlElement
                va = OxmlElement('w:vertAlign')
                va.set(qn('w:val'), 'superscript')
                rPr.append(va)
            else:
                new_run.text = part
            
            # 统一应用字体格式，字号保持原样不上小号
            new_run.font.size = Pt(size)
            new_run.font.name = font_west
            new_run.font.bold = is_bold
            new_run.font.color.rgb = RGBColor(0, 0, 0)
            r = new_run._element.get_or_add_rPr()
            r_fonts = r.get_or_add_rFonts()
            r_fonts.set(qn('w:eastAsia'), font_east)
    else:
        # 没有暗号的段落，按原逻辑处理
        for run in paragraph.runs:
            run.font.size = Pt(size)
            run.font.bold = is_bold
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.font.name = font_west
            r = run._element.get_or_add_rPr()
            r_fonts = r.get_or_add_rFonts()
            r_fonts.set(qn('w:eastAsia'), font_east)

def apply_custom_styles(docx_path):
    """
    Post-processing the Word document to inject specific fonts and spacing.
    """
    doc = Document(docx_path)
    
    for para in doc.paragraphs:
        style_name = para.style.name
        
        if style_name == 'Heading 1':
            # H1: 22pt (Level 2), Heiti, Center, Bold, 1.5x
            set_format(para, '黑体', 'Times New Roman', 22, True, WD_ALIGN_PARAGRAPH.CENTER, is_title=True)
        elif style_name == 'Heading 2':
            # H2: 15pt (Small 3), SimSun, Left, Bold, 1.5x
            set_format(para, '宋体', 'Times New Roman', 15, True, WD_ALIGN_PARAGRAPH.LEFT, is_title=True)
        elif style_name == 'Heading 3':
            # H3: 12pt (Small 4), SimSun, Left, Bold, 1.5x
            set_format(para, '宋体', 'Times New Roman', 12, True, WD_ALIGN_PARAGRAPH.LEFT, is_title=True)
        else:
            # Body: 10.5pt (Level 5), SimSun, Left, 1.25x
            set_format(para, '宋体', 'Times New Roman', 10.5, False, WD_ALIGN_PARAGRAPH.LEFT, is_title=False)
            
    doc.save(docx_path)

def convert_file(input_path):
    output_path = input_path.rsplit('.', 1)[0] + '.docx'
    
    # --- 新增预处理：读取 md，将 <sup> 替换为暗号 ---
    with open(input_path, 'r', encoding='utf-8') as f:
        content = f.read()
    # 将 <sup>...</sup> 替换为 {{SUP:内容}}
    processed_content = re.sub(r'<sup>(.*?)</sup>', r'{{SUP:\1}}', content)
    
    temp_md = input_path + ".temp.md"
    with open(temp_md, 'w', encoding='utf-8') as f:
        f.write(processed_content)
    
    extra_args = ['--mathjax', '--from=markdown+tex_math_dollars']
    
    try:
        # 使用处理后的临时文件转换
        pypandoc.convert_file(temp_md, 'docx', outputfile=output_path, extra_args=extra_args)
        apply_custom_styles(output_path)
        print(f"  [SUCCESS] Created & Formatted: {os.path.basename(output_path)}")
    finally:
        if os.path.exists(temp_md):
            os.remove(temp_md) # 删除临时文件

def main():
    try:
        version = pypandoc.get_pandoc_version()
        print(f"--- [SYSTEM] Pandoc Environment: OK (Version {version}) ---")
    except OSError:
        print("\n[!] ERROR: Pandoc not found in System Path.")
        return

    print("\n" + "="*50)
    print("      MARKDOWN TO WORD CONVERTER (v3.2 PRO)")
    print("      Styles: Songti/TNR | H:1.5x | Body:1.25x")
    print("="*50)
    print("Choose your conversion mode:")
    print("  1. SINGLE FILE (Convert 1 specific .md file)")
    print("  2. ALL FILES   (Convert every .md file in a folder)")
    print("-" * 50)
    
    choice = input("Enter selection (1 or 2): ").strip()

    if choice == '1':
        print("\n[MODE: Single File]")
        path = input(">>> Please drag/paste the [FILE] path here: ").strip('"').strip()
        if os.path.isfile(path) and path.lower().endswith('.md'):
            convert_file(path)
        else:
            print("Invalid input! Make sure it is a valid .md file path.")

    elif choice == '2':
        print("\n[MODE: All Files in Folder]")
        folder = input(">>> Please drag/paste the [FOLDER] path here: ").strip('"').strip()
        if os.path.isdir(folder):
            files = [os.path.join(folder, f) for f in os.listdir(folder) if f.lower().endswith('.md')]
            if not files:
                print("No Markdown (.md) files found in this directory.")
            else:
                print(f"Processing {len(files)} files...")
                for f in files:
                    convert_file(f)
        else:
            print("Invalid directory path!")
    else:
        print("Selection error! Please run the script again and enter 1 or 2.")

    print("\n" + "="*50)
    print("Task completed. Have a great day!")
    input("Press [Enter] to close this window...")

if __name__ == "__main__":
    main()